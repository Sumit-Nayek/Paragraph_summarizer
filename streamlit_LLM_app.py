# import os
# import streamlit as st
# from dotenv import load_dotenv
# from openai import OpenAI
# from pypdf import PdfReader
# from docx import Document

# # Load local environment variables from .env
# load_dotenv()

# # --- 1. RESOLVE API KEY ---
# nvidia_api_key = st.secrets.get("NVIDIA_API_KEY") or os.getenv("NVIDIA_API_KEY")

# # --- 2. STREAMLIT UI CONFIGURATION ---
# st.set_page_config(
#     page_title="AI Paragraph Summarizer (NVIDIA NIM)",
#     page_icon="⚡",
#     layout="centered"
# )

# st.title("⚡ AI Paragraph Summarizer")
# st.caption("Powered by NVIDIA NIM API & Llama Models")

# # --- 3. HELPER FUNCTIONS FOR FILE PARSING & METRICS ---
# def extract_text_from_pdf(file) -> str:
#     """Extracts text from an uploaded PDF file using pypdf."""
#     reader = PdfReader(file)
#     extracted_text = []
#     for page in reader.pages:
#         text = page.extract_text()
#         if text:
#             extracted_text.append(text)
#     return "\n".join(extracted_text)

# def extract_text_from_docx(file) -> str:
#     """Extracts text from an uploaded Word (.docx) file using python-docx."""
#     doc = Document(file)
#     extracted_text = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
#     return "\n".join(extracted_text)

# def extract_text_from_txt(file) -> str:
#     """Extracts text from an uploaded plain text (.txt) file."""
#     return file.read().decode("utf-8")

# def display_analytics(original_text: str, summary_text: str):
#     """Calculates and displays summary metrics in columns."""
#     orig_words = len(original_text.split())
#     sum_words = len(summary_text.split())
    
#     if orig_words > 0:
#         reduction = round(((orig_words - sum_words) / orig_words) * 100, 1)
#         reduction = max(0.0, reduction) # Ensure non-negative
#         time_saved_min = round((orig_words - sum_words) / 200, 1) # Avg reading speed: 200 wpm
#         time_saved_min = max(0.0, time_saved_min)
#     else:
#         reduction = 0.0
#         time_saved_min = 0.0

#     st.markdown("---")
#     st.subheader("📊 Summary Analytics")
#     col1, col2, col3, col4 = st.columns(4)
    
#     col1.metric("Original Words", f"{orig_words}")
#     col2.metric("Summary Words", f"{sum_words}")
#     col3.metric("Reduction", f"{reduction}%")
#     col4.metric("Time Saved", f"~{time_saved_min} min")


# # --- 4. SIDEBAR CONTROLS ---
# st.sidebar.header("🚀 Speed & Model Settings")

# # Model Mode Selector (Fast vs. Deep)
# model_mode = st.sidebar.radio(
#     "Select Inference Mode:",
#     ["⚡ Fast Mode (Llama 3.1 8B)", "🧠 Deep Mode (Llama 3.3 70B)"],
#     help="Fast Mode gives sub-second responses; Deep Mode provides higher reasoning and detail."
# )

# if "Fast Mode" in model_mode:
#     selected_model = "meta/llama-3.1-8b-instruct"
# else:
#     selected_model = "meta/llama-3.3-70b-instruct"

# # Customization Settings
# st.sidebar.markdown("---")
# st.sidebar.header("🎯 Summary Style & Language")

# summary_style = st.sidebar.selectbox(
#     "Summary Format / Style:",
#     [
#         "Executive Bullet Points",
#         "TL;DR One-Liner",
#         "ELI5 (Explain Like I'm 5)",
#         "Structured Key Takeaways"
#     ]
# )

# target_language = st.sidebar.selectbox(
#     "Target Language:",
#     ["English", "Spanish", "French", "German", "Hindi", "Mandarin Chinese"]
# )

# # Generation Hyperparameters
# st.sidebar.markdown("---")
# st.sidebar.header("🎛️ Hyperparameters")
# temperature = st.sidebar.slider("Creativity (Temperature)", 0.0, 1.0, 0.2, 0.1)
# max_tokens = st.sidebar.slider("Max Summary Tokens", 50, 1000, 300, 50)

# # Check for API Key presence
# if not nvidia_api_key:
#     st.error("⚠️ NVIDIA API Key not found. Please set `NVIDIA_API_KEY` in Streamlit Secrets or your `.env` file.")
#     st.stop()

# # --- 5. INITIALIZE OPENAI CLIENT POINTING TO NVIDIA NIM ---
# client = OpenAI(
#     base_url="https://integrate.api.nvidia.com/v1",
#     api_key=nvidia_api_key
# )

# # --- 6. DYNAMIC STREAMING GENERATOR FUNCTION ---
# def generate_summary_stream(text: str, model: str, temp: float, tokens: int, style: str, language: str):
#     """Generator function that builds dynamic system prompts and streams tokens from NVIDIA NIM."""
    
#     style_prompts = {
#         "Executive Bullet Points": "Summarize the text using clear, professional executive bullet points.",
#         "TL;DR One-Liner": "Provide a single, impactful one-sentence summary (TL;DR).",
#         "ELI5 (Explain Like I'm 5)": "Explain and summarize the main concepts in very simple terms as if explaining to a 5-year-old.",
#         "Structured Key Takeaways": "Provide a brief 2-sentence overview followed by a section of key takeaways."
#     }

#     system_instruction = (
#         f"You are an expert multilingual summarizer. "
#         f"{style_prompts.get(style, '')} "
#         f"IMPORTANT: The final summary MUST be written in {language}."
#     )

#     response_stream = client.chat.completions.create(
#         model=model,
#         messages=[
#             {"role": "system", "content": system_instruction},
#             {"role": "user", "content": f"Please summarize the following text:\n\n{text}"}
#         ],
#         temperature=temp,
#         max_tokens=tokens,
#         stream=True
#     )

#     for chunk in response_stream:
#         if chunk.choices and chunk.choices[0].delta.content:
#             yield chunk.choices[0].delta.content


# # --- 7. MAIN INPUT (FILE UPLOAD OR MANUAL TEXT) ---
# st.subheader("1. Input Document or Text")

# uploaded_file = st.file_uploader(
#     "Upload a document (.pdf, .docx, .txt):",
#     type=["pdf", "docx", "txt"]
# )

# extracted_text = ""

# if uploaded_file is not None:
#     try:
#         if uploaded_file.name.endswith(".pdf"):
#             extracted_text = extract_text_from_pdf(uploaded_file)
#         elif uploaded_file.name.endswith(".docx"):
#             extracted_text = extract_text_from_docx(uploaded_file)
#         elif uploaded_file.name.endswith(".txt"):
#             extracted_text = extract_text_from_txt(uploaded_file)
            
#         st.success(f"Successfully loaded '{uploaded_file.name}'!")
#     except Exception as e:
#         st.error(f"Error reading file: {e}")

# input_text = st.text_area(
#     "Or edit/paste your text below:",
#     value=extracted_text,
#     height=220,
#     placeholder="Enter long text or upload a file above..."
# )

# # --- 8. GENERATION LOGIC & ANALYTICS ---
# st.subheader("2. Summary Output")

# if st.button("Generate Real-Time Summary", type="primary"):
#     if not input_text.strip():
#         st.warning("Please upload a valid file or enter text before generating a summary.")
#     else:
#         try:
#             # st.write_stream streams tokens and returns the concatenated full string when done
#             full_summary = st.write_stream(
#                 generate_summary_stream(
#                     text=input_text,
#                     model=selected_model,
#                     temp=temperature,
#                     tokens=max_tokens,
#                     style=summary_style,
#                     language=target_language
#                 )
#             )
            
#             # Save to session_state so it persists across button re-renders
#             st.session_state["current_summary"] = full_summary
#             st.session_state["current_input"] = input_text

#         except Exception as e:
#             st.error(f"An error occurred while calling NVIDIA NIM API: {e}")

# # Display Analytics and Download Button if a summary exists in session_state
# if "current_summary" in st.session_state and st.session_state["current_summary"]:
#     full_summary = st.session_state["current_summary"]
#     current_input = st.session_state["current_input"]
    
#     # Display Analytics Dashboard
#     display_analytics(original_text=current_input, summary_text=full_summary)
    
#     # One-Click Download Button
#     st.markdown("### 📥 Export Summary")
#     st.download_button(
#         label="Download Summary (.txt)",
#         data=full_summary,
#         file_name="summary_output.txt",
#         mime="text/plain",
#         type="secondary"
#     )
import streamlit as st
import requests
from openai import OpenAI
from pypdf import PdfReader
from docx import Document

# --- 1. STREAMLIT UI CONFIGURATION ---
st.set_page_config(
    page_title="AI Content Suite",
    page_icon="⚡",
    layout="centered"
)

st.title("⚡ AI Content Suite")
st.caption("Summarize text with NVIDIA NIM and detect AI-generated content via Hugging Face.")

# --- 2. ZERO-TRUST SECURITY: SIDEBAR API KEY INPUTS ---
st.sidebar.header("🔑 API Credentials")
st.sidebar.markdown("Keys are kept securely in temporary memory and are never saved.")

nvidia_api_key = st.sidebar.text_input("NVIDIA API Key (for Summarizer)", type="password")
hf_token = st.sidebar.text_input("Hugging Face Token (for AI Detector)", type="password")

st.sidebar.markdown("---")


# --- 3. HELPER FUNCTIONS FOR FILE PARSING & METRICS ---
def extract_text_from_pdf(file) -> str:
    reader = PdfReader(file)
    return "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])

def extract_text_from_docx(file) -> str:
    doc = Document(file)
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

def extract_text_from_txt(file) -> str:
    return file.read().decode("utf-8")

def display_analytics(original_text: str, summary_text: str):
    orig_words = len(original_text.split())
    sum_words = len(summary_text.split())
    
    reduction = max(0.0, round(((orig_words - sum_words) / orig_words) * 100, 1)) if orig_words > 0 else 0.0
    time_saved_min = max(0.0, round((orig_words - sum_words) / 200, 1)) if orig_words > 0 else 0.0

    st.markdown("---")
    st.subheader("📊 Summary Analytics")
    col1, col2, col3, col4 = st.columns(4)
    col1.metric("Original Words", f"{orig_words}")
    col2.metric("Summary Words", f"{sum_words}")
    col3.metric("Reduction", f"{reduction}%")
    col4.metric("Time Saved", f"~{time_saved_min} min")


# --- 4. STREAMING GENERATOR FOR NVIDIA NIM ---
def generate_summary_stream(text: str, model: str, temp: float, tokens: int, style: str, language: str):
    client = OpenAI(
        base_url="https://integrate.api.nvidia.com/v1",
        api_key=nvidia_api_key
    )
    
    style_prompts = {
        "Executive Bullet Points": "Summarize the text using clear, professional executive bullet points.",
        "TL;DR One-Liner": "Provide a single, impactful one-sentence summary (TL;DR).",
        "ELI5 (Explain Like I'm 5)": "Explain and summarize the main concepts in very simple terms as if explaining to a 5-year-old.",
        "Structured Key Takeaways": "Provide a brief 2-sentence overview followed by a section of key takeaways."
    }

    system_instruction = (
        f"You are an expert multilingual summarizer. {style_prompts.get(style, '')} "
        f"IMPORTANT: The final summary MUST be written in {language}."
    )

    response_stream = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": system_instruction},
            {"role": "user", "content": f"Please summarize the following text:\n\n{text}"}
        ],
        temperature=temp,
        max_tokens=tokens,
        stream=True
    )

    for chunk in response_stream:
        if chunk.choices and chunk.choices[0].delta.content:
            yield chunk.choices[0].delta.content


# --- 5. TABS SETUP ---
tab_summary, tab_ai_detect = st.tabs(["📝 Summarizer", "🤖 AI Content Detector"])


# ==========================================
# TAB 1: PARAGRAPH SUMMARIZER
# ==========================================
with tab_summary:
    st.header("Paragraph Summarizer")
    
    # Settings exclusively for the summarizer
    with st.expander("⚙️ Summarizer Settings"):
        model_mode = st.radio("Inference Mode:", ["⚡ Fast Mode (Llama 3.1 8B)", "🧠 Deep Mode (Llama 3.3 70B)"])
        selected_model = "meta/llama-3.1-8b-instruct" if "Fast Mode" in model_mode else "meta/llama-3.3-70b-instruct"
        
        col_style, col_lang = st.columns(2)
        summary_style = col_style.selectbox("Format:", ["Executive Bullet Points", "TL;DR One-Liner", "ELI5 (Explain Like I'm 5)", "Structured Key Takeaways"])
        target_language = col_lang.selectbox("Language:", ["English", "Spanish", "French", "German", "Hindi", "Mandarin Chinese"])
        
        temperature = st.slider("Creativity (Temperature)", 0.0, 1.0, 0.2, 0.1)
        max_tokens = st.slider("Max Summary Tokens", 50, 1000, 300, 50)

    # Input Area
    uploaded_file = st.file_uploader("Upload a document (.pdf, .docx, .txt):", type=["pdf", "docx", "txt"], key="sum_file")
    extracted_text = ""
    
    if uploaded_file:
        try:
            if uploaded_file.name.endswith(".pdf"): extracted_text = extract_text_from_pdf(uploaded_file)
            elif uploaded_file.name.endswith(".docx"): extracted_text = extract_text_from_docx(uploaded_file)
            elif uploaded_file.name.endswith(".txt"): extracted_text = extract_text_from_txt(uploaded_file)
            st.success("File loaded successfully!")
        except Exception as e:
            st.error(f"Error reading file: {e}")

    input_text = st.text_area("Or edit/paste your text below:", value=extracted_text, height=200)

    # Generate Button
    if st.button("Generate Real-Time Summary", type="primary"):
        if not nvidia_api_key:
            st.error("⚠️ Please enter your NVIDIA API Key in the sidebar first.")
        elif not input_text.strip():
            st.warning("Please upload a file or enter text to summarize.")
        else:
            try:
                full_summary = st.write_stream(
                    generate_summary_stream(input_text, selected_model, temperature, max_tokens, summary_style, target_language)
                )
                st.session_state["current_summary"] = full_summary
                st.session_state["current_input"] = input_text
            except Exception as e:
                st.error(f"API Error: {e}")

    # Analytics & Download
    if st.session_state.get("current_summary"):
        display_analytics(st.session_state["current_input"], st.session_state["current_summary"])
        st.download_button("📥 Download Summary (.txt)", data=st.session_state["current_summary"], file_name="summary.txt", mime="text/plain")


# ==========================================
# TAB 2: AI CONTENT DETECTOR
# ==========================================
with tab_ai_detect:
    st.header("AI Content Detector")
    st.caption("Detects whether a text was likely written by a human or generated by AI.")
    
    detect_input = st.text_area("Paste text to analyze:", height=250, key="detect_input", placeholder="Paste the paragraph you want to verify here...")
    
    if st.button("Scan Text", type="primary"):
        if not hf_token:
            st.error("⚠️ Please enter your Hugging Face Token in the sidebar first.")
        elif not detect_input.strip():
            st.warning("Please enter some text to scan.")
        else:
            with st.spinner("Scanning for AI patterns..."):
                # Define Hugging Face Inference API details
                API_URL = "https://api-inference.huggingface.co/models/openai-community/roberta-base-openai-detector"
                headers = {"Authorization": f"Bearer {hf_token}"}
                payload = {"inputs": detect_input}
                
                try:
                    # Make the HTTP POST request
                    response = requests.post(API_URL, headers=headers, json=payload)
                    
                    if response.status_code == 200:
                        results = response.json()
                        
                        # Hugging Face returns a list of lists. Example: [[{'label': 'Fake', 'score': 0.99}, ...]]
                        if isinstance(results, list) and len(results) > 0:
                            st.subheader("Detection Results:")
                            col1, col2 = st.columns(2)
                            
                            for classification in results[0]:
                                label = classification.get("label")
                                score = classification.get("score", 0) * 100
                                
                                if label == "Fake" or label == "AI":
                                    col1.metric("🤖 Probability it is AI", f"{score:.2f}%")
                                else:
                                    col2.metric("🧑‍💻 Probability it is Human", f"{score:.2f}%")
                        else:
                            st.warning("No detection scores returned.")
                            
                    elif response.status_code == 503:
                        st.warning("The detection model is currently loading on Hugging Face's servers (Cold Start). Please wait 20 seconds and click 'Scan Text' again.")
                    else:
                        st.error(f"Error {response.status_code}: {response.text}")
                        
                except Exception as e:
                    st.error(f"Detection failed: {e}")